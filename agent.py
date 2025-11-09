


from dotenv import load_dotenv
import streamlit as st
from typing import List, Optional, Dict, Any
import docx
import io
import json
from pathlib import Path
import pandas as pd
from PyPDF2 import PdfReader
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langchain_community.vectorstores import FAISS
from langchain_core.runnables import RunnablePassthrough
from langchain_core.documents import Document
from langchain_core.messages import AIMessage, HumanMessage
import langid
import re

# Charger les variables d'environnement
load_dotenv()

# Initialiser l'état de la session
def init_session_state():
    """Initialise l'état de la session Streamlit"""
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "retriever" not in st.session_state:
        st.session_state.retriever = None
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "user_language" not in st.session_state:
        st.session_state.user_language = "fr"
    if "patient_data" not in st.session_state:
        st.session_state.patient_data = None
    if "file_processed" not in st.session_state:
        st.session_state.file_processed = False

# Définir le modèle LLM
@st.cache_resource
def get_llm():
    """Retourne l'instance du modèle de langage"""
    return ChatGroq(model="llama-3.1-8b-instant")

# Définir les embeddings
@st.cache_resource
def get_embeddings():
    """Retourne l'instance des embeddings"""
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Template de prompt amélioré
MEDICAL_ANALYSIS_TEMPLATE = """
Vous êtes un assistant médical expert analysant les ordonnances et documents de santé des patients.

CONTEXTE MÉDICAL DU PATIENT:
{context}

QUESTION DE L'UTILISATEUR: {question}

INSTRUCTIONS:
- Si l'utilisateur envoie une salutation (bonjour, hello, hi, salut, etc.), répondez de manière courtoise et professionnelle sans analyser la santé
- Si l'utilisateur demande une analyse médicale, procédez à l'analyse complète
- Analysez l'ordonnance dans son ensemble sans mentionner de variables spécifiques comme "produit" ou "diagnostic"
- Basez votre analyse sur les médicaments prescrits et les informations médicales disponibles
- Identifiez les conditions médicales probables basées sur les traitements prescrits
- Signalez les interactions médicamenteuses potentielles
- Fournissez des conseils personnalisés adaptés à la situation du patient
- Recommandez toujours de consulter un médecin pour un diagnostic précis

FORMAT DE RÉPONSE POUR L'ANALYSE MÉDICALE:

🩺 **Analyse des Traitements Prescrits:**
- Liste des médicaments et leur utilisation probable
- Conditions médicales identifiées basées sur les prescriptions

🔍 **Évaluation de l'État de Santé:**
- Analyse détaillée basée sur l'ensemble des traitements
- Facteurs de risque potentiels

⚠️ **Alerte Interactions Médicamenteuses:**
- Interactions potentielles à surveiller
- Précautions particulières

💡 **Recommandations:**
- Conseils personnalisés pour le patient
- Suivi médical recommandé

📋 **Important:** Cette analyse est basée sur l'ordonnance et ne remplace pas une consultation médicale. Consultez toujours un médecin pour un diagnostic précis.

Langue de réponse: {language}
"""

GREETING_RESPONSES = {
    'fr': "Bonjour ! Je suis votre assistant médical. Je peux vous aider à analyser votre ordonnance et vos traitements. Comment puis-je vous assister aujourd'hui ?",
    'en': "Hello! I'm your medical assistant. I can help you analyze your prescription and treatments. How can I assist you today?",
    'es': "¡Hola! Soy su asistente médico. Puedo ayudarle a analizar su receta y tratamientos. ¿Cómo puedo ayudarle hoy?",
    'de': "Hallo! Ich bin Ihr medizinischer Assistent. Ich kann Ihnen bei der Analyse Ihres Rezepts und Ihrer Behandlungen helfen. Wie kann ich Ihnen heute behilflich sein?",
    'it': "Ciao! Sono il tuo assistente medico. Posso aiutarti ad analizzare la tua prescrizione e i tuoi trattamenti. Come posso assisterti oggi?",
    'ar': "مرحباً! أنا مساعدك الطبي. يمكنني مساعدتك في تحليل وصفتك الطبية وعلاجاتك. كيف يمكنني مساعدتك اليوم؟"
}

def detect_language(text: str) -> str:
    """Détecte la langue du texte"""
    try:
        lang, confidence = langid.classify(text)
        return lang if lang in ['fr', 'en', 'es', 'de', 'it', 'ar'] else 'fr'
    except Exception:
        return 'fr'

def is_greeting(message: str) -> bool:
    """Détecte si le message est une salutation"""
    greetings_patterns = [
        # Français
        r'\b(?:bonjour|salut|coucou|hello|hi|hey|bonsoir)\b',
        # Anglais
        r'\b(?:hello|hi|hey|greetings|good morning|good afternoon)\b',
        # Espagnol
        r'\b(?:hola|buenos días|buenas tardes)\b',
        # Allemand
        r'\b(?:hallo|guten tag|guten morgen)\b',
        # Italien
        r'\b(?:ciao|buongiorno|buonasera)\b',
        # Arabe
        r'\b(?:مرحبا|السلام عليكم|اهلا)\b'
    ]
    
    message_lower = message.lower().strip()
    for pattern in greetings_patterns:
        if re.search(pattern, message_lower, re.IGNORECASE):
            return True
    return False

def extract_text_from_file(uploaded_file) -> str:
    """Extrait le texte de différents types de fichiers"""
    if uploaded_file is None:
        return ""
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_extension == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_extension == 'txt':
            return extract_text_from_txt(uploaded_file)
        elif file_extension in ['doc', 'docx']:
            return extract_text_from_docx(uploaded_file)
        elif file_extension == 'csv':
            return extract_text_from_csv(uploaded_file)
        else:
            raise ValueError(f"Type de fichier non supporté: {file_extension}")
    except Exception as e:
        raise Exception(f"Erreur lors du traitement du fichier {file_extension.upper()}: {str(e)}")

def extract_text_from_pdf(uploaded_file) -> str:
    """Extrait le texte d'un fichier PDF"""
    pdf_reader = PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text.strip()

def extract_text_from_txt(uploaded_file) -> str:
    """Extrait le texte d'un fichier TXT"""
    return uploaded_file.getvalue().decode('utf-8').strip()

def extract_text_from_docx(uploaded_file) -> str:
    """Extrait le texte d'un fichier DOCX"""
    doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
    return "\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()

def extract_text_from_csv(uploaded_file) -> str:
    """Extrait et formate le texte d'un fichier CSV"""
    try:
        df = pd.read_csv(uploaded_file)
    except Exception:
        uploaded_file.seek(0)
        df = pd.read_csv(uploaded_file, encoding='latin1', sep=';')
    
    # Normaliser les noms de colonnes
    df.columns = df.columns.str.lower().str.strip().str.replace(' ', '_')
    
    # Construire une représentation textuelle naturelle
    text_representation = "Ordonnance médicale - Liste des traitements:\n\n"
    
    for _, row in df.iterrows():
        line_items = []
        for col_name, value in row.items():
            if pd.notna(value):
                line_items.append(f"{col_name}: {value}")
        
        if line_items:
            text_representation += " | ".join(line_items) + "\n"
    
    return text_representation.strip()

def split_text(text: str) -> List[Document]:
    """Divise le texte en chunks pour le traitement"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [Document(page_content=chunk) for chunk in chunks]

def save_chunks_with_embeddings(documents: List[Document], embeddings, folder_name="chunks_embeddings"):
    """Sauvegarde les chunks avec leurs embeddings"""
    Path(folder_name).mkdir(parents=True, exist_ok=True)
    
    for i, doc in enumerate(documents):
        embedding = embeddings.embed_query(doc.page_content)
        chunk_data = {
            "text": doc.page_content,
            "embedding": embedding
        }
        
        file_path = Path(folder_name) / f"chunk_{i}.json"
        with open(file_path, "w", encoding='utf-8') as f:
            json.dump(chunk_data, f, indent=4, ensure_ascii=False)

def create_retriever(documents: List[Document]):
    """Crée le retriever pour la recherche de documents"""
    embeddings = get_embeddings()
    save_chunks_with_embeddings(documents, embeddings)
    
    vector_store = FAISS.from_documents(
        documents=documents,
        embedding=embeddings,
    )
    vector_store.save_local("vector_store")
    
    return vector_store.as_retriever(search_kwargs={"k": 3})

def create_conversation_chain(retriever):
    """Crée la chaîne de conversation"""
    llm = get_llm()

    def format_context(inputs):
        """Formate le contexte pour la génération de réponse"""
        question = inputs["question"]
        st.session_state.user_language = detect_language(question)
        
        # Si c'est une salutation, on ne cherche pas dans les documents
        if is_greeting(question):
            return {
                **inputs,
                "context": "",
                "language": st.session_state.user_language,
                "is_greeting": True
            }
        
        # Pour les questions médicales, on cherche dans les documents
        docs = retriever.invoke(question)
        return {
            **inputs,
            "context": "\n\n".join([doc.page_content for doc in docs]),
            "language": st.session_state.user_language,
            "is_greeting": False
        }

    def generate_response(inputs):
        """Génère la réponse appropriée"""
        if inputs.get("is_greeting", False):
            return GREETING_RESPONSES.get(inputs["language"], GREETING_RESPONSES['fr'])
        
        # Utiliser le template médical pour les autres questions
        prompt = ChatPromptTemplate.from_template(MEDICAL_ANALYSIS_TEMPLATE)
        chain = prompt | llm | StrOutputParser()
        
        return chain.invoke({
            "context": inputs["context"],
            "question": inputs["question"],
            "language": inputs["language"]
        })

    def conversation_chain(question):
        """Chaîne de conversation principale"""
        formatted_inputs = format_context({"question": question})
        return generate_response(formatted_inputs)
    
    return conversation_chain

def display_patient_data(uploaded_file):
    """Affiche les données du patient de manière structurée"""
    if uploaded_file.name.endswith('.csv'):
        uploaded_file.seek(0)
        try:
            df = pd.read_csv(uploaded_file)
        except Exception:
            uploaded_file.seek(0)
            df = pd.read_csv(uploaded_file, encoding='latin1', sep=';')
        
        # Normaliser les noms de colonnes
        df.columns = df.columns.str.lower().str.strip().str.replace(' ', '_')
        
        if not df.empty:
            st.session_state.patient_data = df
            return True
    return False

def main():
    """Fonction principale de l'application"""
    st.set_page_config(
        page_title="Assistant Médical - Analyse d'Ordonnances",
        page_icon="🩺",
        layout="wide"
    )
    
    st.title("🩺 Assistant Médical - Analyse des Ordonnances")
    st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .sidebar .sidebar-content {
        background-color: #f8f9fa;
    }
    </style>
    """, unsafe_allow_html=True)
    
    init_session_state()

    # Sidebar pour l'upload des fichiers
    with st.sidebar:
        st.header("📄 Upload des Documents Médicaux")
        st.markdown("---")
        
        uploaded_file = st.file_uploader(
            "Téléchargez l'ordonnance ou le document médical du patient", 
            type=["csv", "pdf", "txt", "doc", "docx"],
            help="Formats supportés: CSV, PDF, TXT, DOC, DOCX"
        )
        
        if uploaded_file and not st.session_state.file_processed:
            if st.button("🔍 Analyser le Document", type="primary", use_container_width=True):
                with st.spinner("Analyse en cours..."):
                    try:
                        extracted_text = extract_text_from_file(uploaded_file)
                        
                        if not extracted_text:
                            st.error("❌ Aucune donnée n'a pu être extraite du document.")
                            return
                        
                        # Traiter les documents
                        documents = split_text(extracted_text)
                        st.session_state.retriever = create_retriever(documents)
                        st.session_state.conversation = create_conversation_chain(st.session_state.retriever)
                        st.session_state.file_processed = True
                        
                        # Afficher les données si c'est un CSV
                        if uploaded_file.name.endswith('.csv'):
                            if display_patient_data(uploaded_file):
                                st.success("✅ Données médicales importées avec succès!")
                                
                                st.subheader("Traitements du Patient:")
                                st.dataframe(
                                    st.session_state.patient_data,
                                    use_container_width=True,
                                    hide_index=True
                                )
                        else:
                            st.success("✅ Document médical analysé avec succès!")
                            
                    except Exception as e:
                        st.error(f"❌ Erreur lors de l'analyse: {str(e)}")
                        return

        if st.session_state.file_processed:
            if st.button("🔄 Nouveau Document", use_container_width=True):
                st.session_state.file_processed = False
                st.session_state.chat_history = []
                st.session_state.retriever = None
                st.session_state.conversation = None
                st.session_state.patient_data = None
                st.rerun()

    # Zone de chat principale
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("💬 Dialogue avec l'Assistant Médical")
        
        if st.session_state.conversation is None:
            st.info("""
            👋 **Bienvenue dans votre assistant médical !**
            
            Pour commencer :
            1. 📄 Téléchargez une ordonnance ou un document médical dans la sidebar
            2. 🔍 Cliquez sur 'Analyser le Document' 
            3. 💬 Posez vos questions sur les traitements et la santé du patient
            
            **Formats supportés:** CSV, PDF, DOC, DOCX, TXT
            """)
            
            # Exemple de questions
            with st.expander("💡 Exemples de questions que vous pouvez poser"):
                st.markdown("""
                - **Quels sont les traitements prescrits dans cette ordonnance ?**
                - **Quelles conditions médicales peuvent justifier ces médicaments ?**
                - **Y a-t-il des interactions médicamenteuses à surveiller ?**
                - **Pouvez-vous expliquer le rôle de chaque médicament ?**
                - **Quels conseils donneriez-vous à ce patient ?**
                """)

        else:
            # Afficher l'historique de chat
            for message in st.session_state.chat_history:
                with st.chat_message("Human" if isinstance(message, HumanMessage) else "AI"):
                    st.markdown(message.content)

            # Input de chat
            query = st.chat_input("Posez votre question sur les traitements du patient...")
            if query:
                # Ajouter le message utilisateur à l'historique
                st.session_state.chat_history.append(HumanMessage(content=query))
                
                # Afficher le message utilisateur
                with st.chat_message("Human"):
                    st.markdown(query)

                # Générer et afficher la réponse
                with st.chat_message("AI"):
                    response_container = st.empty()
                    
                    try:
                        response = st.session_state.conversation(query)
                        response_container.markdown(response)
                        st.session_state.chat_history.append(AIMessage(content=response))
                    except Exception as e:
                        error_msg = f"⚠️ Désolé, une erreur s'est produite. Veuillez réessayer."
                        response_container.markdown(error_msg)
                        st.session_state.chat_history.append(AIMessage(content=error_msg))

    with col2:
        st.subheader("📋 Informations Importantes")
        st.markdown("""
        <div style='background-color: #f0f2f6; padding: 20px; border-radius: 10px; border-left: 4px solid #1f77b4;'>
        <h4 style='color: #1f77b4; margin-top: 0;'>⚠️ Avertissement Médical</h4>
        <p style='font-size: 0.9rem; margin-bottom: 0;'>
        Cet assistant fournit des informations générales basées sur l'analyse des ordonnances. 
        Il ne remplace pas l'avis d'un professionnel de santé. 
        <strong>Consultez toujours un médecin</strong> pour un diagnostic et un traitement personnalisé.
        </p>
        </div>
        """, unsafe_allow_html=True)
        
        if st.session_state.patient_data is not None:
            st.markdown("---")
            st.subheader("📊 Résumé des Traitements")
            if len(st.session_state.patient_data.columns) >= 2:
                col1_name = st.session_state.patient_data.columns[0]
                col2_name = st.session_state.patient_data.columns[1]
                
                st.metric(
                    label="Nombre de traitements",
                    value=len(st.session_state.patient_data)
                )

if __name__ == "__main__":
    main()